import os
import shutil
import json
from docx import Document

class TranslateService:

    file_name = None
    t2s = None
    s2t = None
    
    # translate docx from simplified to traditional
    def translate_docx(self: object, file_path: str, action: str):
        """
        0. Load json file
        1. Copy docx file and rename docx file
        2. Replace text in docx
        """
        # Step 0: Load json file
        self.load_json_files()
        
        # Step 1: Copy docx file and rename docx file
        self.file_name = os.path.basename(file_path)
        doc_file_path = os.path.join(os.getcwd(), self.file_name)
        shutil.copy(file_path, doc_file_path)

        outputFileName = self.file_name.replace('.docx', '')
        if action == 's2t':
            outputFileName = outputFileName+'_TW'
        else:
            outputFileName = outputFileName+'_CN'
        modified_docx_path = outputFileName+'.docx'

        # Step 2: Replace the text in DOCX
        self.replace_text(modified_docx_path, action)

        print(action)

    # load json files
    def load_json_files(self):
        self.t2s = self.load_json_file_as_map(os.path.join(self.get_app_directory(), '../dictionaries', 't2s-char.min.json'))
        self.s2t = self.load_json_file_as_map(os.path.join(self.get_app_directory(), '../dictionaries', 's2t-char.min.json'))
        return

    def get_app_directory(self):
        """Return the application's base directory."""
        return os.path.dirname(os.path.abspath(__file__))

    def load_json_file_as_map(self,file_path):
        with open(file_path, 'r', encoding="utf-8") as file:
            data = json.load(file)
        return data

    def translate(self, action, content):
        # Replace text based on the translation mapping
        if action == 't2s':
            for traditional, simplified in self.t2s.items():
                content = content.replace(traditional, simplified)
        elif action == 's2t':
            for simplified, traditional in self.s2t.items():
                content = content.replace(simplified, traditional)
        return content

    def replace_text(self, docx_path, action):
        doc = Document(docx_path)

        for i, paragraph in enumerate(doc.paragraphs):

            """ replace text in paragraph runs (normal body) """
            for j, run in enumerate(paragraph.runs):
                if run.text != '':
                    # original_text = run.text
                    run.text = self.translate(action, run.text)
                    # translated_text = run.text
                    # print(f"Run {i}_{j}: Original Text: '{original_text}', Translated Text: '{translated_text}'.")

            """ replace text in headings (heading) """
            if paragraph.style.name.startswith("Heading"):
                paragraph.text = self.translate(action, paragraph.text)

        """ replace text in tables """
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph.text = self.translate(action, paragraph.text)

        # TODO: replace text in header and footer are currently not supported, header and footer are read only        

        doc.save(docx_path)